"""
Module 1: OCR + Resume Text Extraction
Handles PDF, DOCX, and image-based resumes
"""

import os
import re
import io

def extract_text_from_file(filepath):
    """Extract raw text from PDF, DOCX, or image file."""
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == '.pdf':
        return _extract_from_pdf(filepath)
    elif ext in ['.docx', '.doc']:
        return _extract_from_docx(filepath)
    elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
        return _extract_from_image(filepath)
    else:
        return ""

def _extract_from_pdf(filepath):
    """Extract text from PDF - tries pdfplumber first, then OCR fallback."""
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"pdfplumber failed: {e}")

    # If no text extracted (scanned PDF), try OCR
    if len(text.strip()) < 50:
        try:
            from pdf2image import convert_from_path
            import pytesseract
            # On Windows, set tesseract path if needed:
            # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            images = convert_from_path(filepath)
            for img in images:
                text += pytesseract.image_to_string(img) + "\n"
        except Exception as e:
            print(f"OCR fallback failed: {e}")

    return text.strip()

def _extract_from_docx(filepath):
    """Extract text from DOCX file."""
    try:
        import docx
        doc = docx.Document(filepath)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return ""

def _extract_from_image(filepath):
    """Extract text from image using Tesseract OCR."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(filepath)
        return pytesseract.image_to_string(img)
    except Exception as e:
        print(f"Image OCR failed: {e}")
        return ""

def clean_text(text):
    """Clean and normalize extracted text."""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep useful ones
    text = re.sub(r'[^\w\s\.\,\-\+\@\#\/\(\)\:\;]', ' ', text)
    # Normalize newlines
    text = text.replace('|', '\n').replace('•', '\n•')
    return text.strip()
